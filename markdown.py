import streamlit as st
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown2 import markdown
from bs4 import BeautifulSoup

# Function to clean the Markdown content
def clean_markdown(md_content):
    # Remove callouts formatted like [!sometext]
    md_content = re.sub(r'\[\!.*?\]', '', md_content)
    # Remove all double brackets [[ or ]]
    md_content = md_content.replace('[[', '').replace(']]', '')
    return md_content

# Function to parse Markdown to Docx
def parse_markdown_to_docx(md_content, docx_file):
    # Clean the markdown content
    md_content = clean_markdown(md_content)
    
    # Convert markdown to HTML
    html_content = markdown(md_content)

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Markdown to Word Conversion', 0)

    # Function to add paragraph with optional bold and italic text
    def add_paragraph(text, style=None, bold=False, italic=False):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if style:
            paragraph.style = style
        return paragraph

    # Parse HTML elements and add them to the Word document
    for element in soup.children:
        if element.name == 'p':
            add_paragraph(element.text)
        elif element.name == 'h1':
            add_paragraph(element.text, 'Heading 1')
        elif element.name == 'h2':
            add_paragraph(element.text, 'Heading 2')
        elif element.name == 'h3':
            add_paragraph(element.text, 'Heading 3')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                add_paragraph('• ' + li.text)
        elif element.name == 'blockquote':
            quote = add_paragraph(element.text, style='Quote')
            quote.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif element.name == 'strong':
            add_paragraph(element.text, bold=True)
        elif element.name == 'em':
            add_paragraph(element.text, italic=True)
        elif element.name == 'code':
            add_paragraph(element.text, style='Code')
        else:
            add_paragraph(element.text)

    # Save the document
    doc.save(docx_file)

# Streamlit application
st.title("Markdown to Word Converter")

# Options for input
option = st.selectbox("Choose the input method", ["Upload Markdown File", "Paste Markdown Text"])

md_content = ""

if option == "Upload Markdown File":
    uploaded_file = st.file_uploader("Choose a Markdown file", type="md")
    if uploaded_file is not None:
        md_content = uploaded_file.read().decode('utf-8')
elif option == "Paste Markdown Text":
    md_content = st.text_area("Paste your Markdown text here")

if md_content:
    # Define the output file path
    output_path = "converted_document.docx"
    
    # Convert the Markdown to Docx
    parse_markdown_to_docx(md_content, output_path)
    
    # Provide a download button for the converted file
    with open(output_path, "rb") as file:
        st.download_button(
            label="Download Word Document",
            data=file,
            file_name="converted_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
